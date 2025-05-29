#!/usr/bin/env python3
"""
AI Agent Hackathon - Healthcare Treatment Analysis System
Complete implementation using AWS Strands SDK
"""

import os
import json
import logging
import boto3
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
import asyncio
from dataclasses import dataclass

# Strands imports
from strands import Agent, tool
from strands.models import BedrockModel

# Additional imports for file handling
import PyPDF2
import docx
from bs4 import BeautifulSoup
import mammoth
import pandas as pd

# Configure logging
logging.getLogger("strands").setLevel(logging.DEBUG)
logging.basicConfig(
    format="%(levelname)s | %(name)s | %(message)s",
    handlers=[logging.StreamHandler()]
)

@dataclass
class TreatmentData:
    """Data structure for treatment information"""
    treatment_id: str
    source: str
    raw_content: str
    structured_content: str = ""
    risk_assessment: str = ""
    revenue_analysis: str = ""

@dataclass
class ProcessingResults:
    """Container for all processing results"""
    treatments: List[TreatmentData]
    final_report: str = ""
    approval_status: str = "pending"

class HealthcareAgentSystem:
    """Main system orchestrator for the healthcare agent hackathon"""
    
    def __init__(self):
        # Initialize AWS Bedrock session
        self.session = boto3.Session(
            aws_access_key_id='ASIAUFJ05VRONE JHNEP7',
            aws_secret_access_key='5FNHUatrDaqx97YKRAVPZ0a2gFgCUW3217RFN9MV',
            aws_session_token='IQ0Jb3JpZ2lux2VjEMn//////////wEaCXVZLWVhc3QtMSJGMEQCIG1TUFYEK5bxNkvX036wTbnjAEJCXOUF 1v2X450pBfJYAIA14t61Qitsis4jN1SBPIWNIWZ6NHn6',
            region_name='us-west-2'
        )
        
        # Create Bedrock model
        self.bedrock_model = BedrockModel(
            model_id="anthropic.claude-3-5-sonnet-20241022-v2:0",
            boto_session=self.session
        )
        
        # Initialize agents
        self.research_agent = None
        self.documenting_agent = None
        self.risk_assessment_agent = None
        self.revenue_identification_agent = None
        self.emailing_agent = None
        
        # Initialize data storage
        self.processing_results = ProcessingResults(treatments=[])
        
        # Setup output directory
        self.output_dir = Path("./hackathon_output")
        self.output_dir.mkdir(exist_ok=True)
        
        self._initialize_agents()
    
    def _initialize_agents(self):
        """Initialize all specialized agents"""
        self.research_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_file_reader_tool(),
                self._create_html_parser_tool(),
                self._create_pdf_parser_tool(),
                self._create_docx_parser_tool()
            ]
        )
        
        self.documenting_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_document_structuring_tool(),
                self._create_report_generator_tool()
            ]
        )
        
        self.risk_assessment_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_risk_analysis_tool(),
                self._create_medical_knowledge_tool()
            ]
        )
        
        self.revenue_identification_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_revenue_analysis_tool(),
                self._create_customer_segmentation_tool()
            ]
        )
        
        self.emailing_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_document_aggregation_tool(),
                self._create_word_export_tool()
            ]
        )

    # ==================== RESEARCH AGENT TOOLS ====================
    
    def _create_file_reader_tool(self):
        @tool
        def read_file(file_path: str) -> Dict[str, Any]:
            """
            Read various file formats from local filesystem.
            Supports .html, .pdf, .docx, .txt files.
            
            Args:
                file_path: Path to the file to read
                
            Returns:
                Dictionary containing file content and metadata
            """
            try:
                path = Path(file_path)
                if not path.exists():
                    return {"error": f"File not found: {file_path}"}
                
                file_extension = path.suffix.lower()
                content = ""
                
                if file_extension == '.html':
                    with open(path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif file_extension == '.txt':
                    with open(path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif file_extension == '.pdf':
                    with open(path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        content = ""
                        for page in pdf_reader.pages:
                            content += page.extract_text() + "\n"
                elif file_extension == '.docx':
                    doc = docx.Document(path)
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                else:
                    return {"error": f"Unsupported file format: {file_extension}"}
                
                return {
                    "file_path": str(path),
                    "file_type": file_extension,
                    "content": content,
                    "size": len(content),
                    "timestamp": datetime.now().isoformat()
                }
                
            except Exception as e:
                return {"error": f"Error reading file {file_path}: {str(e)}"}
        
        return read_file
    
    def _create_html_parser_tool(self):
        @tool
        def parse_html_content(html_content: str, target_treatments: List[str] = None) -> Dict[str, Any]:
            """
            Parse HTML content and extract treatment information.
            
            Args:
                html_content: Raw HTML content
                target_treatments: List of treatment IDs to look for
                
            Returns:
                Parsed treatment data
            """
            try:
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Extract text content
                text_content = soup.get_text(separator=' ', strip=True)
                
                # Look for treatment patterns
                treatments_found = []
                if target_treatments:
                    for treatment in target_treatments:
                        if treatment.lower() in text_content.lower():
                            treatments_found.append(treatment)
                
                # Extract structured data (tables, lists, etc.)
                tables = []
                for table in soup.find_all('table'):
                    table_data = []
                    for row in table.find_all('tr'):
                        row_data = [cell.get_text(strip=True) for cell in row.find_all(['td', 'th'])]
                        table_data.append(row_data)
                    tables.append(table_data)
                
                lists = []
                for ul in soup.find_all(['ul', 'ol']):
                    list_items = [li.get_text(strip=True) for li in ul.find_all('li')]
                    lists.append(list_items)
                
                return {
                    "text_content": text_content,
                    "treatments_found": treatments_found,
                    "tables": tables,
                    "lists": lists,
                    "title": soup.title.string if soup.title else "No title",
                    "headings": [h.get_text(strip=True) for h in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])]
                }
                
            except Exception as e:
                return {"error": f"Error parsing HTML: {str(e)}"}
        
        return parse_html_content
    
    def _create_pdf_parser_tool(self):
        @tool
        def extract_pdf_structured_data(pdf_content: str) -> Dict[str, Any]:
            """
            Extract structured information from PDF content.
            
            Args:
                pdf_content: Raw PDF text content
                
            Returns:
                Structured data extracted from PDF
            """
            try:
                # Split content into sections
                lines = pdf_content.split('\n')
                sections = []
                current_section = ""
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Check if line might be a heading (simple heuristic)
                    if len(line) < 100 and any(keyword in line.lower() for keyword in 
                                             ['treatment', 'procedure', 'therapy', 'diagnosis', 'symptom']):
                        if current_section:
                            sections.append(current_section)
                        current_section = line + "\n"
                    else:
                        current_section += line + "\n"
                
                if current_section:
                    sections.append(current_section)
                
                # Extract key medical terms
                medical_terms = []
                medical_keywords = ['diagnosis', 'treatment', 'therapy', 'procedure', 'medication', 
                                  'dosage', 'side effect', 'contraindication', 'efficacy']
                
                for line in lines:
                    for keyword in medical_keywords:
                        if keyword in line.lower():
                            medical_terms.append(line.strip())
                
                return {
                    "sections": sections,
                    "medical_terms": medical_terms,
                    "total_lines": len(lines),
                    "content_length": len(pdf_content)
                }
                
            except Exception as e:
                return {"error": f"Error processing PDF content: {str(e)}"}
        
        return extract_pdf_structured_data
    
    def _create_docx_parser_tool(self):
        @tool
        def parse_docx_content(docx_content: str) -> Dict[str, Any]:
            """
            Parse DOCX content for treatment information.
            
            Args:
                docx_content: Text content from DOCX file
                
            Returns:
                Parsed treatment data
            """
            try:
                paragraphs = [p.strip() for p in docx_content.split('\n') if p.strip()]
                
                # Identify potential treatment sections
                treatment_sections = []
                for i, paragraph in enumerate(paragraphs):
                    if any(keyword in paragraph.lower() for keyword in 
                          ['treatment', 'therapy', 'procedure', 'intervention']):
                        # Get context (previous and next paragraphs)
                        context = []
                        start = max(0, i-2)
                        end = min(len(paragraphs), i+3)
                        for j in range(start, end):
                            context.append(paragraphs[j])
                        treatment_sections.append({
                            'heading': paragraph,
                            'context': context
                        })
                
                return {
                    "paragraphs": paragraphs,
                    "treatment_sections": treatment_sections,
                    "paragraph_count": len(paragraphs)
                }
                
            except Exception as e:
                return {"error": f"Error parsing DOCX content: {str(e)}"}
        
        return parse_docx_content

    # ==================== DOCUMENTING AGENT TOOLS ====================
    
    def _create_document_structuring_tool(self):
        @tool
        def structure_treatment_data(raw_data: Dict[str, Any], treatment_id: str) -> Dict[str, Any]:
            """
            Structure raw treatment data into organized format.
            
            Args:
                raw_data: Raw extracted data from various sources
                treatment_id: ID of the treatment being processed
                
            Returns:
                Structured treatment document
            """
            try:
                structured_doc = {
                    "treatment_id": treatment_id,
                    "timestamp": datetime.now().isoformat(),
                    "data_sources": [],
                    "treatment_overview": "",
                    "clinical_details": {},
                    "regulatory_info": {},
                    "cost_information": {},
                    "effectiveness_data": {}
                }
                
                # Process content based on source type
                if isinstance(raw_data, dict):
                    if "content" in raw_data:
                        content = raw_data["content"]
                        
                        # Extract treatment overview
                        if "treatment" in content.lower():
                            lines = content.split('\n')
                            overview_lines = []
                            for line in lines:
                                if any(keyword in line.lower() for keyword in 
                                      ['treatment', 'therapy', 'procedure', 'overview']):
                                    overview_lines.append(line.strip())
                            structured_doc["treatment_overview"] = ' '.join(overview_lines[:3])
                        
                        # Extract clinical details
                        clinical_keywords = ['dosage', 'administration', 'duration', 'frequency']
                        for keyword in clinical_keywords:
                            if keyword in content.lower():
                                # Find sentences containing the keyword
                                sentences = content.split('.')
                                relevant_sentences = [s.strip() for s in sentences if keyword in s.lower()]
                                structured_doc["clinical_details"][keyword] = relevant_sentences[:2]
                        
                        # Extract effectiveness data
                        effectiveness_keywords = ['efficacy', 'success rate', 'outcome', 'result']
                        for keyword in effectiveness_keywords:
                            if keyword in content.lower():
                                sentences = content.split('.')
                                relevant_sentences = [s.strip() for s in sentences if keyword in s.lower()]
                                structured_doc["effectiveness_data"][keyword] = relevant_sentences[:2]
                        
                        structured_doc["data_sources"].append({
                            "file_path": raw_data.get("file_path", "unknown"),
                            "file_type": raw_data.get("file_type", "unknown"),
                            "processing_timestamp": datetime.now().isoformat()
                        })
                
                return structured_doc
                
            except Exception as e:
                return {"error": f"Error structuring treatment data: {str(e)}"}
        
        return structure_treatment_data
    
    def _create_report_generator_tool(self):
        @tool
        def generate_combined_report(structured_treatments: List[Dict[str, Any]]) -> str:
            """
            Generate a combined report from multiple structured treatments.
            
            Args:
                structured_treatments: List of structured treatment documents
                
            Returns:
                Combined report as formatted text
            """
            try:
                report_sections = []
                
                # Header
                report_sections.append("="*60)
                report_sections.append("HEALTHCARE TREATMENT ANALYSIS REPORT")
                report_sections.append("="*60)
                report_sections.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                report_sections.append(f"Total Treatments Analyzed: {len(structured_treatments)}")
                report_sections.append("")
                
                # Executive Summary
                report_sections.append("EXECUTIVE SUMMARY")
                report_sections.append("-" * 20)
                treatment_ids = [t.get("treatment_id", "Unknown") for t in structured_treatments]
                report_sections.append(f"Treatments Covered: {', '.join(treatment_ids)}")
                report_sections.append("")
                
                # Individual Treatment Analysis
                for i, treatment in enumerate(structured_treatments, 1):
                    report_sections.append(f"TREATMENT {i}: {treatment.get('treatment_id', 'Unknown')}")
                    report_sections.$append("-" * 40)
                    
                    # Treatment Overview
                    overview = treatment.get("treatment_overview", "No overview available")
                    report_sections.append(f"Overview: {overview}")
                    report_sections.append("")
                    
                    # Clinical Details
                    clinical = treatment.get("clinical_details", {})
                    if clinical:
                        report_sections.append("Clinical Details:")
                        for key, value in clinical.items():
                            if isinstance(value, list):
                                report_sections.append(f"  - {key.title()}: {'; '.join(value)}")
                            else:
                                report_sections.append(f"  - {key.title()}: {value}")
                        report_sections.append("")
                    
                    # Effectiveness Data
                    effectiveness = treatment.get("effectiveness_data", {})
                    if effectiveness:
                        report_sections.append("Effectiveness Data:")
                        for key, value in effectiveness.items():
                            if isinstance(value, list):
                                report_sections.append(f"  - {key.title()}: {'; '.join(value)}")
                            else:
                                report_sections.append(f"  - {key.title()}: {value}")
                        report_sections.append("")
                    
                    # Data Sources
                    sources = treatment.get("data_sources", [])
                    if sources:
                        report_sections.append("Data Sources:")
                        for source in sources:
                            report_sections.append(f"  - {source.get('file_path', 'Unknown')}")
                        report_sections.append("")
                    
                    report_sections.append("")
                
                return "\n".join(report_sections)
                
            except Exception as e:
                return f"Error generating combined report: {str(e)}"
        
        return generate_combined_report

    # ==================== RISK ASSESSMENT AGENT TOOLS ====================
    
    def _create_risk_analysis_tool(self):
        @tool
        def analyze_treatment_risks(treatment_data: Dict[str, Any]) -> Dict[str, Any]:
            """
            Analyze risks associated with a treatment.
            
            Args:
                treatment_data: Structured treatment data
                
            Returns:
                Risk analysis results
            """
            try:
                risk_analysis = {
                    "treatment_id": treatment_data.get("treatment_id", "Unknown"),
                    "overall_risk_score": 0,  # Scale of 1-10
                    "medical_risks": [],
                    "financial_risks": [],
                    "regulatory_risks": [],
                    "recommendations": [],
                    "analysis_timestamp": datetime.now().isoformat()
                }
                
                # Analyze medical risks
                content = str(treatment_data)
                medical_risk_keywords = {
                    'high': ['contraindication', 'severe', 'fatal', 'death', 'emergency'],
                    'medium': ['side effect', 'adverse', 'caution', 'monitor', 'warning'],
                    'low': ['mild', 'temporary', 'reversible', 'minor']
                }
                
                medical_risk_score = 1
                for risk_level, keywords in medical_risk_keywords.items():
                    for keyword in keywords:
                        if keyword in content.lower():
                            risk_analysis["medical_risks"].append({
                                "keyword": keyword,
                                "level": risk_level,
                                "context": f"Found in treatment data: {keyword}"
                            })
                            if risk_level == 'high':
                                medical_risk_score += 3
                            elif risk_level == 'medium':
                                medical_risk_score += 2
                            elif risk_level == 'low':
                                medical_risk_score += 1
                
                # Analyze financial risks
                financial_keywords = ['expensive', 'cost', 'insurance', 'coverage', 'reimbursement']
                financial_risk_score = 1
                for keyword in financial_keywords:
                    if keyword in content.lower():
                        risk_analysis["financial_risks"].append({
                            "factor": keyword,
                            "impact": "medium",
                            "description": f"Financial consideration: {keyword}"
                        })
                        financial_risk_score += 1
                
                # Analyze regulatory risks
                regulatory_keywords = ['fda', 'approval', 'trial', 'experimental', 'investigational']
                regulatory_risk_score = 1
                for keyword in regulatory_keywords:
                    if keyword in content.lower():
                        risk_analysis["regulatory_risks"].append({
                            "factor": keyword,
                            "status": "requires_review",
                            "description": f"Regulatory consideration: {keyword}"
                        })
                        regulatory_risk_score += 1
                
                # Calculate overall risk score
                risk_analysis["overall_risk_score"] = min(10, 
                    (medical_risk_score + financial_risk_score + regulatory_risk_score) / 3)
                
                # Generate recommendations
                if risk_analysis["overall_risk_score"] >= 7:
                    risk_analysis["recommendations"].append("High risk treatment - requires extensive review")
                    risk_analysis["recommendations"].append("Consider alternative treatments")
                elif risk_analysis["overall_risk_score"] >= 4:
                    risk_analysis["recommendations"].append("Medium risk treatment - proceed with caution")
                    risk_analysis["recommendations"].append("Monitor patient outcomes closely")
                else:
                    risk_analysis["recommendations"].append("Low risk treatment - suitable for coverage")
                
                return risk_analysis
                
            except Exception as e:
                return {"error": f"Error analyzing treatment risks: {str(e)}"}
        
        return analyze_treatment_risks
    
    def _create_medical_knowledge_tool(self):
        @tool
        def consult_medical_knowledge(treatment_id: str, query: str) -> Dict[str, Any]:
            """
            Consult medical knowledge base for additional treatment information.
            
            Args:
                treatment_id: ID of the treatment
                query: Specific medical query
                
            Returns:
                Medical knowledge consultation results
            """
            try:
                # Simulated medical knowledge base
                medical_knowledge = {
                    "treatment_1": {
                        "category": "Cardiovascular",
                        "complexity": "High",
                        "typical_duration": "6-12 months",
                        "success_rate": "85%",
                        "common_complications": ["bleeding", "infection"]
                    },
                    "treatment_2": {
                        "category": "Orthopedic",
                        "complexity": "Medium",
                        "typical_duration": "3-6 months",
                        "success_rate": "90%",
                        "common_complications": ["swelling", "limited mobility"]
                    },
                    "treatment_3": {
                        "category": "Neurological",
                        "complexity": "High",
                        "typical_duration": "12-24 months",
                        "success_rate": "70%",
                        "common_complications": ["cognitive effects", "fatigue"]
                    },
                    "treatment_4": {
                        "category": "Oncological",
                        "complexity": "Very High",
                        "typical_duration": "6-18 months",
                        "success_rate": "75%",
                        "common_complications": ["nausea", "hair loss", "immunosuppression"]
                    },
                    "treatment_7": {
                        "category": "Endocrine",
                        "complexity": "Medium",
                        "typical_duration": "Lifelong",
                        "success_rate": "95%",
                        "common_complications": ["weight changes", "mood changes"]
                    },
                    "treatment_8": {
                        "category": "Respiratory",
                        "complexity": "Medium",
                        "typical_duration": "1-3 months",
                        "success_rate": "88%",
                        "common_complications": ["cough", "throat irritation"]
                    }
                }
                
                knowledge = medical_knowledge.get(treatment_id, {
                    "category": "General",
                    "complexity": "Unknown",
                    "typical_duration": "Variable",
                    "success_rate": "Unknown",
                    "common_complications": []
                })
                
                return {
                    "treatment_id": treatment_id,
                    "query": query,
                    "knowledge_base_info": knowledge,
                    "consultation_timestamp": datetime.now().isoformat(),
                    "confidence_level": "simulated_data"
                }
                
            except Exception as e:
                return {"error": f"Error consulting medical knowledge: {str(e)}"}
        
        return consult_medical_knowledge

    # ==================== REVENUE IDENTIFICATION AGENT TOOLS ====================
    
    def _create_revenue_analysis_tool(self):
        @tool
        def analyze_revenue_opportunities(treatment_data: Dict[str, Any], risk_data: Dict[str, Any]) -> Dict[str, Any]:
            """
            Analyze revenue opportunities for a treatment.
            
            Args:
                treatment_data: Structured treatment data
                risk_data: Risk analysis results
                
            Returns:
                Revenue opportunity analysis
            """
            try:
                revenue_analysis = {
                    "treatment_id": treatment_data.get("treatment_id", "Unknown"),
                    "market_potential": {},
                    "pricing_strategy": {},
                    "customer_segments": {},
                    "revenue_projections": {},
                    "business_case": {},
                    "analysis_timestamp": datetime.now().isoformat()
                }
                
                # Analyze market potential based on treatment complexity and success rate
                risk_score = risk_data.get("overall_risk_score", 5)
                
                if risk_score <= 3:
                    market_potential = "High"
                    base_price_multiplier = 1.2
                elif risk_score <= 6:
                    market_potential = "Medium"
                    base_price_multiplier = 1.0
                else:
                    market_potential = "Low"
                    base_price_multiplier = 0.8
                
                revenue_analysis["market_potential"] = {
                    "level": market_potential,
                    "risk_adjusted_score": 10 - risk_score,
                    "market_readiness": "ready" if risk_score <= 5 else "requires_evaluation"
                }
                
                # Pricing strategy
                base_coverage_amount = 10000  # Base amount in USD
                suggested_coverage = base_coverage_amount * base_price_multiplier
                
                revenue_analysis["pricing_strategy"] = {
                    "base_coverage_amount": base_coverage_amount,
                    "risk_multiplier": base_price_multiplier,
                    "suggested_coverage_limit": suggested_coverage,
                    "premium_adjustment": f"{int((base_price_multiplier - 1) * 100)}%"
                }
                
                # Customer segmentation
                customer_segments = {
                    "new_customers": {
                        "target_demographic": "Age 25-65, health-conscious",
                        "estimated_size": 50000 if market_potential == "High" else 25000,
                        "conversion_rate": "15%" if risk_score <= 4 else "8%"
                    },
                    "existing_customers": {
                        "upgrade_potential": "High" if risk_score <= 4 else "Medium",
                        "estimated_candidates": 20000 if market_potential == "High" else 10000,
                        "upsell_rate": "25%" if risk_score <= 4 else "15%"
                    }
                }
                
                revenue_analysis["customer_segments"] = customer_segments
                
                # Revenue projections
                new_customer_revenue = (customer_segments["new_customers"]["estimated_size"] * 
                                      (0.15 if risk_score <= 4 else 0.08) * suggested_coverage * 0.1)
                existing_customer_revenue = (customer_segments["existing_customers"]["estimated_candidates"] * 
                                           (0.25 if risk_score <= 4 else 0.15) * suggested_coverage * 0.05)
                
                total_projected_revenue = new_customer_revenue + existing_customer_revenue
                
                revenue_analysis["revenue_projections"] = {
                    "new_customer_revenue": new_customer_revenue,
                    "existing_customer_revenue": existing_customer_revenue,
                    "total_annual_projection": total_projected_revenue,
                    "roi_estimate": f"{int(total_projected_revenue / suggested_coverage * 100)}%"
                }
                
                # Business case
                business_case_strength = "Strong" if total_projected_revenue > 1000000 else "Moderate" if total_projected_revenue > 500000 else "Weak"
                
                revenue_analysis["business_case"] = {
                    "strength": business_case_strength,
                    "key_benefits": [
                        f"Market potential: {market_potential}",
                        f"Projected revenue: ${total_projected_revenue:,.2f}",
                        f"Risk level: {'Acceptable' if risk_score <= 5 else 'Requires mitigation'}"
                    ],
                    "recommendations": [
                        "Proceed with coverage" if business_case_strength == "Strong" else "Conduct further analysis",
                        "Target new customer acquisition" if new_customer_revenue > existing_customer_revenue else "Focus on existing customer upsell"
                    ]
                }
                
                return revenue_analysis
                
            except Exception as e:
                return {"error": f"Error analyzing revenue opportunities: {str(e)}"}
        
        return analyze_revenue_opportunities
    
    def _create_customer_segmentation_tool(self):
        @tool
        def segment_customers(treatment_analysis: Dict[str, Any]) -> Dict[str, Any]:
            """
            Perform detailed customer segmentation analysis.
            
            Args:
                treatment_analysis: Complete treatment analysis data
                
            Returns:
                Customer segmentation results
            """
            try:
                segmentation = {
                    "treatment_id": treatment_analysis.get("treatment_id", "Unknown"),
                    "segments": {},
                    "targeting_strategy": {},
                    "engagement_recommendations": {},
                    "segmentation_timestamp": datetime.now().isoformat()
                }
                
                # Define customer segments
                segments = {
                    "premium_seekers": {
                        "description": "Customers seeking comprehensive coverage",
                        "characteristics": ["High income", "Risk averse", "Values premium service"],
                        "size_estimate": "15%",
                        "targeting_priority": "High",
                        "messaging": "Comprehensive protection for peace of mind"
                    },
                    "cost_conscious": {
                        "description": "Price-sensitive customers",
                        "characteristics": ["Budget focused", "Compares options", "Values transparency"],
                        "size_estimate": "45%",
                        "targeting_priority": "Medium",
                        "messaging": "Affordable coverage without compromising quality"
                    },
                    "health_optimizers": {
                        "description": "Proactive health management customers",
                        "characteristics": ["Preventive care focused", "Tech-savvy", "Data-driven"],
                        "size_estimate": "25%",
                        "targeting_priority": "High",
                        "messaging": "Advanced treatment options for optimal health outcomes"
                    },
                    "traditional_seekers": {
                        "description": "Customers preferring established treatments",
                        "characteristics": ["Conservative", "Trusts proven methods", "Values stability"],
                        "size_estimate": "15%",
                        "targeting_priority": "Medium",
                        "messaging": "Time-tested treatments with proven results"
                    }
                }
                
                segmentation["segments"] = segments
                
                # Targeting strategy
                segmentation["targeting_strategy"] = {
                    "primary_segments": ["premium_seekers", "health_optimizers"],
                    "secondary_segments": ["cost_conscious", "traditional_seekers"],
                    "channel_preferences": {
                        "premium_seekers": ["direct_sales", "specialist_referrals"],
                        "cost_conscious": ["online_marketing", "comparison_sites"],
                        "health_optimizers": ["digital_channels", "health_apps"],
                        "traditional_seekers": ["traditional_media", "doctor_referrals"]
                    }
                }
                
                # Engagement recommendations
                segmentation["engagement_recommendations"] = {
                    "premium_seekers": {
                        "approach": "Personalized consultation",
                        "content": "Detailed treatment analysis and premium benefits",
                        "timing": "Immediate follow-up"
                    },
                    "cost_conscious": {
                        "approach": "Value-based messaging",
                        "content": "Cost-benefit analysis and coverage options",
                        "timing": "Regular updates"
                    },
                    "health_optimizers": {
                        "approach": "Data-driven engagement",
                        "content": "Treatment effectiveness metrics and outcomes",
                        "timing": "Continuous engagement"
                    },
                    "traditional_seekers": {
                        "approach": "Trust-building communication",
                        "content": "Success stories and expert endorsements",
                        "timing": "Gradual introduction"
                    }
                }
                
                return segmentation
                
            except Exception as e:
                return {"error": f"Error segmenting customers: {str(e)}"}
        
        return segment_customers

    # ==================== EMAILING AGENT TOOLS ====================
    
    def _create_document_aggregation_tool(self):
        @tool
        def aggregate_documents(analysis_results: Dict[str, Any]) -> Dict[str, Any]:
            """
            Aggregate all analysis documents into a single report.
            
            Args:
                analysis_results: Complete analysis results
                
            Returns:
                Aggregated document
            """
            try:
                aggregated_doc = {
                    "timestamp": datetime.now().isoformat(),
                    "treatment_id": analysis_results.get("treatment_id", "Unknown"),
                    "sections": {},
                    "metadata": {}
                }
                
                # Aggregate different analysis sections
                if "risk_analysis" in analysis_results:
                    aggregated_doc["sections"]["risk_analysis"] = {
                        "overall_risk_score": analysis_results["risk_analysis"].get("overall_risk_score", 0),
                        "key_risks": analysis_results["risk_analysis"].get("medical_risks", [])[:3],
                        "recommendations": analysis_results["risk_analysis"].get("recommendations", [])
                    }
                
                if "revenue_analysis" in analysis_results:
                    aggregated_doc["sections"]["revenue_analysis"] = {
                        "market_potential": analysis_results["revenue_analysis"].get("market_potential", {}),
                        "revenue_projections": analysis_results["revenue_analysis"].get("revenue_projections", {}),
                        "business_case": analysis_results["revenue_analysis"].get("business_case", {})
                    }
                
                if "customer_segmentation" in analysis_results:
                    aggregated_doc["sections"]["customer_segmentation"] = {
                        "primary_segments": analysis_results["customer_segmentation"].get("targeting_strategy", {}).get("primary_segments", []),
                        "engagement_recommendations": analysis_results["customer_segmentation"].get("engagement_recommendations", {})
                    }
                
                # Add metadata
                aggregated_doc["metadata"] = {
                    "analysis_version": "1.0",
                    "generated_by": "HealthcareAgentSystem",
                    "data_sources": analysis_results.get("data_sources", []),
                    "processing_time": datetime.now().isoformat()
                }
                
                return aggregated_doc
                
            except Exception as e:
                return {"error": f"Error aggregating documents: {str(e)}"}
        
        return aggregate_documents
    
    def _create_word_export_tool(self):
        @tool
        def export_to_word(aggregated_doc: Dict[str, Any], output_path: str) -> Dict[str, Any]:
            """
            Export aggregated document to Word format.
            
            Args:
                aggregated_doc: Aggregated document
                output_path: Path to save the Word document
                
            Returns:
                Export status
            """
            try:
                # Create a new Word document
                doc = docx.Document()
                
                # Add title
                doc.add_heading('Healthcare Treatment Analysis Report', 0)
                
                # Add timestamp
                doc.add_paragraph(f"Generated: {aggregated_doc['timestamp']}")
                doc.add_paragraph(f"Treatment ID: {aggregated_doc['treatment_id']}")
                doc.add_paragraph("")
                
                # Add sections
                for section_name, section_data in aggregated_doc["sections"].items():
                    doc.add_heading(section_name.replace("_", " ").title(), level=1)
                    
                    if isinstance(section_data, dict):
                        for key, value in section_data.items():
                            if isinstance(value, (list, dict)):
                                doc.add_paragraph(f"{key.replace('_', ' ').title()}:")
                                if isinstance(value, list):
                                    for item in value:
                                        doc.add_paragraph(f"  - {item}", style='List Bullet')
                                else:
                                    for k, v in value.items():
                                        doc.add_paragraph(f"  - {k}: {v}", style='List Bullet')
                            else:
                                doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
                    
                    doc.add_paragraph("")
                
                # Add metadata
                doc.add_heading('Metadata', level=1)
                for key, value in aggregated_doc["metadata"].items():
                    if isinstance(value, list):
                        doc.add_paragraph(f"{key.replace('_', ' ').title()}:")
                        for item in value:
                            doc.add_paragraph(f"  - {item}", style='List Bullet')
                    else:
                        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
                
                # Save the document
                doc.save(output_path)
                
                return {
                    "status": "success",
                    "output_path": output_path,
                    "timestamp": datetime.now().isoformat()
                }
                
            except Exception as e:
                return {"error": f"Error exporting to Word: {str(e)}"}
        
        return export_to_word

    # ==================== MAIN EXECUTION ====================
    
    async def process_treatment(self, treatment_id: str, input_files: List[str]) -> Dict[str, Any]:
        """
        Main method to process a treatment through all agents.
        
        Args:
            treatment_id: ID of the treatment to process
            input_files: List of input file paths
            
        Returns:
            Complete analysis results
        """
        try:
            # Initialize treatment data
            treatment_data = TreatmentData(
                treatment_id=treatment_id,
                source="input_files",
                raw_content=""
            )
            
            # Research phase
            for file_path in input_files:
                file_data = await self.research_agent.run(
                    self._create_file_reader_tool(),
                    file_path=file_path
                )
                
                if "error" not in file_data:
                    treatment_data.raw_content += file_data["content"] + "\n"
            
            # Document structuring phase
            structured_data = await self.documenting_agent.run(
                self._create_document_structuring_tool(),
                raw_data={"content": treatment_data.raw_content},
                treatment_id=treatment_id
            )
            
            treatment_data.structured_content = structured_data
            
            # Risk assessment phase
            risk_analysis = await self.risk_assessment_agent.run(
                self._create_risk_analysis_tool(),
                treatment_data=structured_data
            )
            
            treatment_data.risk_assessment = risk_analysis
            
            # Revenue analysis phase
            revenue_analysis = await self.revenue_identification_agent.run(
                self._create_revenue_analysis_tool(),
                treatment_data=structured_data,
                risk_data=risk_analysis
            )
            
            treatment_data.revenue_analysis = revenue_analysis
            
            # Customer segmentation
            segmentation = await self.revenue_identification_agent.run(
                self._create_customer_segmentation_tool(),
                treatment_analysis={
                    "treatment_id": treatment_id,
                    "risk_analysis": risk_analysis,
                    "revenue_analysis": revenue_analysis
                }
            )
            
            # Document aggregation
            aggregated_doc = await self.emailing_agent.run(
                self._create_document_aggregation_tool(),
                analysis_results={
                    "treatment_id": treatment_id,
                    "risk_analysis": risk_analysis,
                    "revenue_analysis": revenue_analysis,
                    "customer_segmentation": segmentation,
                    "data_sources": [{"file_path": f} for f in input_files]
                }
            )
            
            # Export to Word
            output_path = str(self.output_dir / f"treatment_analysis_{treatment_id}.docx")
            export_result = await self.emailing_agent.run(
                self._create_word_export_tool(),
                aggregated_doc=aggregated_doc,
                output_path=output_path
            )
            
            # Store results
            self.processing_results.treatments.append(treatment_data)
            
            return {
                "status": "success",
                "treatment_id": treatment_id,
                "output_file": output_path,
                "analysis_results": aggregated_doc
            }
            
        except Exception as e:
            return {
                "status": "error",
                "treatment_id": treatment_id,
                "error": str(e)
            }

async def main():
    """Main execution function"""
    # Initialize the system
    system = HealthcareAgentSystem()
    
    # Example treatment processing
    treatment_id = "treatment_1"
    input_files = [
        "path/to/treatment_document.pdf",
        "path/to/clinical_data.docx"
    ]
    
    results = await system.process_treatment(treatment_id, input_files)
    print(json.dumps(results, indent=2))

def run_main():
    """Wrapper function to handle event loop properly"""
    try:
        # Get the current event loop or create a new one
        loop = asyncio.get_event_loop()
    except RuntimeError:
        # If no event loop exists, create a new one
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
    
    try:
        # Run the main async function
        loop.run_until_complete(main())
    finally:
        # Clean up
        loop.close()

if __name__ == "__main__":
    run_main()