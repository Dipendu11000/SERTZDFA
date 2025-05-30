#!/usr/bin/env python3
"""
AI Agent Hackathon - Healthcare Treatment Analysis System
Complete implementation using AWS Strands SDK
"""

import os
import json
import logging
import boto3
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
import asyncio
from dataclasses import dataclass
import sys

# Strands imports
from strands import Agent, tool
from strands.models import BedrockModel

# Additional imports for file handling
import PyPDF2
import docx
from bs4 import BeautifulSoup
import mammoth
import pandas as pd

# Configure logging
logging.getLogger("strands").setLevel(logging.DEBUG)
logging.basicConfig(
    format="%(levelname)s | %(name)s | %(message)s",
    handlers=[logging.StreamHandler()]
)

@dataclass
class TreatmentData:
    """Data structure for treatment information"""
    treatment_id: str
    source: str
    raw_content: str
    structured_content: str = ""
    risk_assessment: str = ""
    revenue_analysis: str = ""

@dataclass
class ProcessingResults:
    """Container for all processing results"""
    treatments: List[TreatmentData]
    final_report: str = ""
    approval_status: str = "pending"

class HealthcareAgentSystem:
    """Main system orchestrator for the healthcare agent hackathon"""
    
    def __init__(self):
        # Initialize AWS Bedrock session
        self.session = boto3.Session(
            aws_access_key_id='ASIAUFJ05VRONE JHNEP7',
            aws_secret_access_key='5FNHUatrDaqx97YKRAVPZ0a2gFgCUW3217RFN9MV',
            aws_session_token='IQ0Jb3JpZ2lux2VjEMn//////////wEaCXVZLWVhc3QtMSJGMEQCIG1TUFYEK5bxNkvX036wTbnjAEJCXOUF 1v2X450pBfJYAIA14t61Qitsis4jN1SBPIWNIWZ6NHn6',
            region_name='us-west-2'
        )
        
        # Create Bedrock model
        self.bedrock_model = BedrockModel(
            model_id="anthropic.claude-3-5-sonnet-20241022-v2:0",
            boto_session=self.session
        )
        # Create a single LLM agent for all LLM calls
        self.llm_agent = Agent(model=self.bedrock_model)
        
        # Initialize agents
        self.research_agent = None
        self.documenting_agent = None
        self.risk_assessment_agent = None
        self.revenue_identification_agent = None
        self.emailing_agent = None
        
        # Initialize data storage
        self.processing_results = ProcessingResults(treatments=[])
        
        # Setup output directory
        self.output_dir = Path("./hackathon_output")
        self.output_dir.mkdir(exist_ok=True)
        
        self._initialize_agents()
    
    def _initialize_agents(self):
        """Initialize all specialized agents"""
        self.research_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_file_reader_tool(),
                self._create_research_agent_tool(),
                self._create_document_structuring_tool(),
                self._create_risk_analysis_tool(),
                self._create_revenue_analysis_tool(),
                self._create_report_generator_tool()
            ]
        )
        
        self.documenting_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_document_structuring_tool(),
                self._create_report_generator_tool()
            ]
        )
        
        self.risk_assessment_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_risk_analysis_tool(),
                self._create_medical_knowledge_tool()
            ]
        )
        
        self.revenue_identification_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_revenue_analysis_tool(),
                self._create_customer_segmentation_tool()
            ]
        )
        
        self.emailing_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_document_aggregation_tool(),
                self._create_word_export_tool()
            ]
        )

    # ==================== RESEARCH AGENT TOOLS ====================
    
    def _create_file_reader_tool(self):
        @tool
        def read_file(file_path: str) -> Dict[str, Any]:
            """
            Read various file formats from local filesystem.
            Supports .html, .pdf, .docx, .txt files.
            
            Args:
                file_path: Path to the file to read
                
            Returns:
                Dictionary containing file content and metadata
            """
            try:
                path = Path(file_path)
                if not path.exists():
                    return {"error": f"File not found: {file_path}"}
                
                file_extension = path.suffix.lower()
                content = ""
                
                if file_extension == '.html':
                    with open(path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif file_extension == '.txt':
                    with open(path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif file_extension == '.pdf':
                    with open(path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        content = ""
                        for page in pdf_reader.pages:
                            content += page.extract_text() + "\n"
                elif file_extension == '.docx':
                    doc = docx.Document(path)
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                else:
                    return {"error": f"Unsupported file format: {file_extension}"}
                
                return {
                    "file_path": str(path),
                    "file_type": file_extension,
                    "content": content,
                    "size": len(content),
                    "timestamp": datetime.now().isoformat()
                }
                
            except Exception as e:
                return {"error": f"Error reading file {file_path}: {str(e)}"}
        
        return read_file
    
    def _create_research_agent_tool(self):
        @tool
        def extract_and_group_treatments(file_contents: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
            """
            Use LLM to extract, group, and merge semantically similar treatments from all input files.
            Returns a list of treatments, each with merged data and a detailed LLM-generated summary.
            """
            try:
                prompt = (
                    "You are an expert medical analyst. Given the following documents, extract all unique treatments, "
                    "group semantically similar treatments, and for each group, merge the information and write a detailed, "
                    "human-readable summary. For each treatment, provide: treatment_id (or assign one if missing), heading, "
                    "detailed description, and list of source files.\n\n"
                    "Documents:\n" +
                    "\n\n".join([f"File: {f['file_path']}\nContent:\n{f['content']}" for f in file_contents])
                )
                result = self.llm_agent(prompt)
                response = result.output if hasattr(result, 'output') else str(result)
                # Expecting a JSON list of treatments
                try:
                    treatments = json.loads(response)
                except Exception:
                    # If not JSON, just return as a single treatment
                    treatments = [{
                        "treatment_id": "unknown",
                        "heading": "All Treatments",
                        "description": response,
                        "source_files": [f["file_path"] for f in file_contents]
                    }]
                return treatments
            except Exception as e:
                return [{"error": f"Error extracting/grouping treatments: {str(e)}"}]
        return extract_and_group_treatments

    # ==================== DOCUMENTING AGENT TOOLS ====================
    
    def _create_document_structuring_tool(self):
        @tool
        def structure_treatment_data_llm(treatments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
            """
            Use LLM to generate a detailed, human-readable, structured document for each treatment.
            """
            try:
                prompt = (
                    "For each treatment below, write a detailed, human-friendly description with headings, context, and explanations.\n\n"
                    f"Treatments:\n{json.dumps(treatments, indent=2)}"
                )
                result = self.llm_agent(prompt)
                response = result.output if hasattr(result, 'output') else str(result)
                try:
                    structured = json.loads(response)
                except Exception:
                    # If not JSON, return as a single block
                    structured = [{"description": response}]
                return structured
            except Exception as e:
                return [{"error": f"Error structuring treatment data: {str(e)}"}]
        return structure_treatment_data_llm

    # ==================== RISK ASSESSMENT AGENT TOOLS ====================
    
    def _create_risk_analysis_tool(self):
        @tool
        def analyze_treatment_risks_llm(treatment: Dict[str, Any]) -> Dict[str, Any]:
            """
            Use LLM to analyze risks, explain parameters, show logic, and generate a detailed, readable risk section.
            """
            try:
                prompt = (
                    "Given the following treatment, perform a risk assessment.\n"
                    "1. List and explain the risk parameters you use.\n"
                    "2. Show how you arrive at the risk score (combine semantic, logical, and mathematical reasoning).\n"
                    "3. Provide a detailed, human-readable risk section with recommendations.\n\n"
                    f"Treatment:\n{json.dumps(treatment, indent=2)}"
                )
                result = self.llm_agent(prompt)
                response = result.output if hasattr(result, 'output') else str(result)
                try:
                    risk = json.loads(response)
                except Exception:
                    risk = {"risk_section": response}
                return risk
            except Exception as e:
                return {"error": f"Error in LLM risk analysis: {str(e)}"}
        return analyze_treatment_risks_llm

    # ==================== REVENUE IDENTIFICATION AGENT TOOLS ====================
    
    def _create_revenue_analysis_tool(self):
        @tool
        def analyze_revenue_opportunities_llm(treatment: Dict[str, Any], risk: Dict[str, Any]) -> Dict[str, Any]:
            """
            Use LLM to analyze revenue opportunities, explain parameters, and generate a business-friendly, readable section.
            """
            try:
                prompt = (
                    "Given the following treatment and its risk analysis, perform a revenue opportunity analysis.\n"
                    "1. List and explain the revenue parameters you use.\n"
                    "2. Show your logic for revenue scoring and segmentation.\n"
                    "3. Provide a detailed, business-friendly revenue section.\n\n"
                    f"Treatment:\n{json.dumps(treatment, indent=2)}\n"
                    f"Risk Analysis:\n{json.dumps(risk, indent=2)}"
                )
                result = self.llm_agent(prompt)
                response = result.output if hasattr(result, 'output') else str(result)
                try:
                    revenue = json.loads(response)
                except Exception:
                    revenue = {"revenue_section": response}
                return revenue
            except Exception as e:
                return {"error": f"Error in LLM revenue analysis: {str(e)}"}
        return analyze_revenue_opportunities_llm

    # ==================== EMAILING AGENT TOOLS ====================
    
    def _create_document_aggregation_tool(self):
        @tool
        def aggregate_documents(analysis_results: Dict[str, Any]) -> Dict[str, Any]:
            """
            Aggregate all analysis documents into a single report.
            
            Args:
                analysis_results: Complete analysis results
                
            Returns:
                Aggregated document
            """
            try:
                aggregated_doc = {
                    "timestamp": datetime.now().isoformat(),
                    "treatment_id": analysis_results.get("treatment_id", "Unknown"),
                    "sections": {},
                    "metadata": {}
                }
                
                # Aggregate different analysis sections
                if "risk_analysis" in analysis_results:
                    aggregated_doc["sections"]["risk_analysis"] = {
                        "overall_risk_score": analysis_results["risk_analysis"].get("overall_risk_score", 0),
                        "key_risks": analysis_results["risk_analysis"].get("medical_risks", [])[:3],
                        "recommendations": analysis_results["risk_analysis"].get("recommendations", [])
                    }
                
                if "revenue_analysis" in analysis_results:
                    aggregated_doc["sections"]["revenue_analysis"] = {
                        "market_potential": analysis_results["revenue_analysis"].get("market_potential", {}),
                        "revenue_projections": analysis_results["revenue_analysis"].get("revenue_projections", {}),
                        "business_case": analysis_results["revenue_analysis"].get("business_case", {})
                    }
                
                if "customer_segmentation" in analysis_results:
                    aggregated_doc["sections"]["customer_segmentation"] = {
                        "primary_segments": analysis_results["customer_segmentation"].get("targeting_strategy", {}).get("primary_segments", []),
                        "engagement_recommendations": analysis_results["customer_segmentation"].get("engagement_recommendations", {})
                    }
                
                # Add metadata
                aggregated_doc["metadata"] = {
                    "analysis_version": "1.0",
                    "generated_by": "HealthcareAgentSystem",
                    "data_sources": analysis_results.get("data_sources", []),
                    "processing_time": datetime.now().isoformat()
                }
                
                return aggregated_doc
                
            except Exception as e:
                return {"error": f"Error aggregating documents: {str(e)}"}
        
        return aggregate_documents
    
    def _create_word_export_tool(self):
        @tool
        def export_to_word(aggregated_doc: list, output_path: str) -> Dict[str, Any]:
            """
            Export a detailed, structured report (JSON list) to a visually appealing Word document.
            """
            try:
                doc = docx.Document()
                doc.add_heading('Healthcare Treatment Analysis Report', 0)
                doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph("")
                for treatment in aggregated_doc:
                    doc.add_heading(treatment.get("treatment_name", "Treatment"), level=1)
                    doc.add_paragraph(treatment.get("description", "No description available."))
                    # Risks Table
                    risks = treatment.get("risks", [])
                    if risks:
                        doc.add_heading("Risk Assessment", level=2)
                        table = doc.add_table(rows=1, cols=3)
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Metric'
                        hdr_cells[1].text = 'Value'
                        hdr_cells[2].text = 'Explanation/Logic'
                        for risk in risks:
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(risk.get("metric", ""))
                            row_cells[1].text = str(risk.get("value", ""))
                            row_cells[2].text = str(risk.get("explanation", ""))
                        doc.add_paragraph("")
                    # Customers Table
                    customers = treatment.get("customers", {})
                    if customers:
                        doc.add_heading("Customer Impact & Segmentation", level=2)
                        table = doc.add_table(rows=1, cols=4)
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Existing Customers'
                        hdr_cells[1].text = 'New Customers'
                        hdr_cells[2].text = 'Segmentation'
                        hdr_cells[3].text = 'Logic/Explanation'
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(customers.get("existing", ""))
                        row_cells[1].text = str(customers.get("new", ""))
                        row_cells[2].text = str(customers.get("segmentation", ""))
                        row_cells[3].text = str(customers.get("logic", ""))
                        doc.add_paragraph("")
                    # Revenue Table
                    revenue = treatment.get("revenue", {})
                    if revenue:
                        doc.add_heading("Revenue & Profitability", level=2)
                        table = doc.add_table(rows=1, cols=4)
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = 'Cost'
                        hdr_cells[1].text = 'Revenue'
                        hdr_cells[2].text = 'Profit'
                        hdr_cells[3].text = 'Logic/Explanation'
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(revenue.get("cost", ""))
                        row_cells[1].text = str(revenue.get("revenue", ""))
                        row_cells[2].text = str(revenue.get("profit", ""))
                        row_cells[3].text = str(revenue.get("logic", ""))
                        doc.add_paragraph("")
                    doc.add_page_break()
                doc.save(output_path)
                return {
                    "status": "success",
                    "output_path": output_path,
                    "timestamp": datetime.now().isoformat()
                }
            except Exception as e:
                return {"error": f"Error exporting to Word: {str(e)}"}
        return export_to_word

    def _create_report_generator_tool(self):
        @tool
        def generate_combined_report_llm(treatments: List[Dict[str, Any]], risks: List[Dict[str, Any]], revenues: List[Dict[str, Any]]) -> list:
            """
            Use LLM to generate a detailed, business-style, treatment-wise report with:
            - Semantic grouping/merging of treatments
            - At least a page of detailed description per treatment
            - Risk metrics with explanations/logic
            - Customer segmentation, revenue, profitability with logic
            - Output as a JSON list: each item has keys: treatment_name, description, risks (list of {metric, value, explanation}), customers (existing, new, segmentation, logic), revenue (cost, revenue, profit, logic)
            """
            try:
                prompt = (
                    "You are a business analyst and medical writer. Given the following treatments, risks, and revenues, generate a detailed, business-style, treatment-wise report. For each treatment, do the following: "
                    "1. Group and merge all semantically similar treatments from all sources.\n"
                    "2. For each treatment, write at least a page of detailed, business-style description (paragraphs, not bullets).\n"
                    "3. For each treatment, provide a table of risk metrics (with value and a detailed explanation/logic for each metric).\n"
                    "4. For each treatment, provide a table of customer impact: estimate how many existing customers will benefit, how many new customers could be acquired (with segmentation, e.g., age group, etc.), and explain the logic.\n"
                    "5. For each treatment, provide a table of cost, revenue, and profitability (with detailed explanation and logic).\n"
                    "6. Output the result as a JSON list, where each item has: treatment_name, description, risks (list of {metric, value, explanation}), customers (existing, new, segmentation, logic), revenue (cost, revenue, profit, logic).\n"
                    "7. Make the report visually appealing and business-friendly.\n\n"
                    f"Treatments:\n{json.dumps(treatments, indent=2)}\n\n"
                    f"Risks:\n{json.dumps(risks, indent=2)}\n\n"
                    f"Revenues:\n{json.dumps(revenues, indent=2)}"
                )
                result = self.llm_agent(prompt)
                response = result.output if hasattr(result, 'output') else str(result)
                try:
                    report_json = json.loads(response)
                except Exception:
                    # fallback: wrap in a single treatment
                    report_json = [{"treatment_name": "All Treatments", "description": response, "risks": [], "customers": {}, "revenue": {}}]
                return report_json
            except Exception as e:
                return [{"treatment_name": "Error", "description": f"Error generating LLM report: {str(e)}", "risks": [], "customers": {}, "revenue": {}}]
        return generate_combined_report_llm

    def _create_medical_knowledge_tool(self):
        @tool
        def consult_medical_knowledge_stub(*args, **kwargs):
            """
            Stub for medical knowledge tool (not used in LLM pipeline, but required for agent initialization).
            """
            return {"info": "Medical knowledge tool is not used in this pipeline."}
        return consult_medical_knowledge_stub

    def _create_customer_segmentation_tool(self):
        @tool
        def customer_segmentation_stub(*args, **kwargs):
            """
            Stub for customer segmentation tool (not used in LLM pipeline, but required for agent initialization).
            """
            return {"info": "Customer segmentation tool is not used in this pipeline."}
        return customer_segmentation_stub

    # ==================== MAIN EXECUTION ====================
    
    async def process_treatment(self, treatment_id: str, input_files: List[str]) -> Dict[str, Any]:
        """
        Main method to process a treatment through all agents.
        
        Args:
            treatment_id: ID of the treatment to process
            input_files: List of input file paths
            
        Returns:
            Complete analysis results
        """
        try:
            print_debug("Reading files...")
            file_reader_tool = self._create_file_reader_tool()
            file_contents = [file_reader_tool(file_path=f) for f in input_files if "error" not in file_reader_tool(file_path=f)]
            print_debug(f"Files read: {len(file_contents)}")
            print_debug("Extracting/grouping treatments with LLM...")
            research_tool = self._create_research_agent_tool()
            treatments = research_tool(file_contents)
            print_debug(f"Treatments extracted: {len(treatments)}")
            print_debug("Generating treatment descriptions with LLM...")
            doc_tool = self._create_document_structuring_tool()
            structured_treatments = doc_tool(treatments)
            print_debug(f"Structured treatments: {len(structured_treatments)}")
            print_debug("Performing risk analysis with LLM...")
            risk_tool = self._create_risk_analysis_tool()
            risks = [risk_tool(t) for t in structured_treatments]
            print_debug(f"Risks analyzed: {len(risks)}")
            print_debug("Performing revenue analysis with LLM...")
            revenue_tool = self._create_revenue_analysis_tool()
            revenues = [revenue_tool(t, r) for t, r in zip(structured_treatments, risks)]
            print_debug(f"Revenues analyzed: {len(revenues)}")
            print_debug("Generating final report with LLM...")
            report_tool = self._create_report_generator_tool()
            final_report = report_tool(structured_treatments, risks, revenues)
            print_debug("Final report generated.")
            # Export the final report to a Word document
            word_export_tool = self._create_word_export_tool()
            output_path = str(self.output_dir / f"treatment_analysis_{treatment_id}.docx")
            export_result = word_export_tool(aggregated_doc=final_report, output_path=output_path)
            print_debug(f"Word document exported to: {output_path}")
            self.processing_results.treatments.append(TreatmentData(
                treatment_id=treatment_id,
                source="input_files",
                raw_content="[LLM processed]",
                structured_content=json.dumps(structured_treatments, indent=2),
                risk_assessment=json.dumps(risks, indent=2),
                revenue_analysis=json.dumps(revenues, indent=2)
            ))
            self.processing_results.final_report = json.dumps(final_report)
            print_debug("Process complete.")
            print(json.dumps(final_report))
            return {
                "status": "success",
                "treatment_id": treatment_id,
                "final_report": json.dumps(final_report),
                "word_output": output_path
            }
        except Exception as e:
            import traceback
            print_debug(f"Exception in process_treatment: {e}")
            traceback.print_exc()
            return {
                "status": "error",
                "treatment_id": treatment_id,
                "error": str(e)
            }

async def main():
    """Main execution function"""
    print_debug("Initializing HealthcareAgentSystem...")
    system = HealthcareAgentSystem()
    treatment_id = "treatment_1"
    input_files = ["competitor.html", "hosp.html"]
    print_debug(f"Processing treatment_id={treatment_id} with files: {input_files}")
    results = await system.process_treatment(treatment_id, input_files)
    print_debug("main() completed.")
    print(results)

def run_main():
    """Wrapper function to handle event loop properly"""
    try:
        # Get the current event loop or create a new one
        loop = asyncio.get_event_loop()
    except RuntimeError:
        # If no event loop exists, create a new one
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
    
    try:
        # Run the main async function
        loop.run_until_complete(main())
    finally:
        # Clean up
        loop.close()

def print_debug(msg):
    print(f"[DEBUG] {msg}", flush=True)

if __name__ == "__main__":
    if "ipykernel" in sys.modules or "IPython" in sys.modules:
        try:
            get_ipython()
            print_debug("Starting main() in Jupyter/IPython...")
            result = await main()
            print_debug("main() completed.")
            print(result)
        except Exception as e:
            print_debug(f"Exception in main(): {e}")
            import traceback
            traceback.print_exc()
    else:
        print_debug("Starting run_main() in script mode...")
        run_main()
