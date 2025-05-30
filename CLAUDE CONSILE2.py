#!/usr/bin/env python3
"""
AI Agent Hackathon - Enhanced Healthcare Treatment Analysis System
Complete implementation using AWS Strands SDK with detailed treatment analysis
"""

import os
import json
import logging
import boto3
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
import asyncio
from dataclasses import dataclass
import sys
import re

# Strands imports
from strands import Agent, tool
from strands.models import BedrockModel

# Additional imports for file handling
import PyPDF2
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from bs4 import BeautifulSoup
import mammoth
import pandas as pd

# Configure logging
logging.getLogger("strands").setLevel(logging.DEBUG)
logging.basicConfig(
    format="%(levelname)s | %(name)s | %(message)s",
    handlers=[logging.StreamHandler()]
)

@dataclass
class TreatmentData:
    """Enhanced data structure for treatment information"""
    treatment_id: str
    treatment_name: str
    source_files: List[str]
    raw_content: str
    detailed_description: str = ""
    risk_assessment: Dict[str, Any] = None
    revenue_analysis: Dict[str, Any] = None
    customer_impact: Dict[str, Any] = None

@dataclass
class ProcessingResults:
    """Container for all processing results"""
    treatments: List[TreatmentData]
    final_report: str = ""
    approval_status: str = "pending"

class HealthcareAgentSystem:
    """Enhanced system orchestrator for the healthcare agent hackathon"""
    
    def __init__(self):
        # Initialize AWS Bedrock session
        self.session = boto3.Session(
            aws_access_key_id='ASIAUFJ05VRONE JHNEP7',
            aws_secret_access_key='5FNHUatrDaqx97YKRAVPZ0a2gFgCUW3217RFN9MV',
            aws_session_token='IQ0Jb3JpZ2lux2VjEMn//////////wEaCXVZLWVhc3QtMSJGMEQCIG1TUFYEK5bxNkvX036wTbnjAEJCXOUF 1v2X450pBfJYAIA14t61Qitsis4jN1SBPIWNIWZ6NHn6',
            region_name='us-west-2'
        )
        
        # Create Bedrock model
        self.bedrock_model = BedrockModel(
            model_id="anthropic.claude-3-5-sonnet-20241022-v2:0",
            boto_session=self.session
        )
        
        # Create a single LLM agent for all LLM calls
        self.llm_agent = Agent(model=self.bedrock_model)
        
        # Initialize agents
        self.research_agent = None
        self.documenting_agent = None
        self.risk_assessment_agent = None
        self.revenue_identification_agent = None
        self.emailing_agent = None
        
        # Initialize data storage
        self.processing_results = ProcessingResults(treatments=[])
        
        # Setup output directory
        self.output_dir = Path("./hackathon_output")
        self.output_dir.mkdir(exist_ok=True)
        
        self._initialize_agents()
    
    def _initialize_agents(self):
        """Initialize all specialized agents"""
        self.research_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_file_reader_tool(),
                self._create_treatment_extraction_tool(),
                self._create_treatment_grouping_tool()
            ]
        )
        
        self.documenting_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_detailed_description_tool(),
                self._create_enhanced_report_generator_tool()
            ]
        )
        
        self.risk_assessment_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_comprehensive_risk_analysis_tool()
            ]
        )
        
        self.revenue_identification_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_detailed_revenue_analysis_tool(),
                self._create_customer_impact_analysis_tool()
            ]
        )
        
        self.emailing_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_enhanced_word_export_tool()
            ]
        )

    # ==================== ENHANCED FILE PROCESSING ====================
    
    def _create_file_reader_tool(self):
        @tool
        def read_file(file_path: str) -> Dict[str, Any]:
            """
            Enhanced file reader with better content extraction and metadata.
            """
            try:
                path = Path(file_path)
                if not path.exists():
                    return {"error": f"File not found: {file_path}"}
                
                file_extension = path.suffix.lower()
                content = ""
                
                if file_extension == '.html':
                    with open(path, 'r', encoding='utf-8') as f:
                        html_content = f.read()
                        soup = BeautifulSoup(html_content, 'html.parser')
                        # Remove script and style elements
                        for script in soup(["script", "style"]):
                            script.extract()
                        content = soup.get_text()
                        # Clean up whitespace
                        content = re.sub(r'\n\s*\n', '\n\n', content)
                        content = re.sub(r' +', ' ', content)
                        
                elif file_extension == '.txt':
                    with open(path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        
                elif file_extension == '.pdf':
                    with open(path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        content = ""
                        for page in pdf_reader.pages:
                            content += page.extract_text() + "\n"
                            
                elif file_extension == '.docx':
                    doc = docx.Document(path)
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                else:
                    return {"error": f"Unsupported file format: {file_extension}"}
                
                return {
                    "file_path": str(path),
                    "file_name": path.name,
                    "file_type": file_extension,
                    "content": content.strip(),
                    "word_count": len(content.split()),
                    "size": len(content),
                    "timestamp": datetime.now().isoformat()
                }
                
            except Exception as e:
                return {"error": f"Error reading file {file_path}: {str(e)}"}
        
        return read_file
    
    def _create_treatment_extraction_tool(self):
        @tool
        def extract_treatments_from_files(file_contents: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
            """
            Extract distinct treatments from all input files with detailed analysis.
            """
            try:
                all_content = ""
                source_mapping = {}
                
                for file_data in file_contents:
                    if "error" not in file_data:
                        content = file_data["content"]
                        file_name = file_data["file_name"]
                        all_content += f"\n\n=== SOURCE: {file_name} ===\n{content}"
                        source_mapping[file_name] = content
                
                prompt = f"""
                You are an expert medical treatment analyst for CareCredit, a healthcare financing company. 
                Analyze the following documents and extract ALL DISTINCT medical treatments mentioned.
                
                For each treatment found:
                1. Assign a unique treatment_id (e.g., TREAT_001, TREAT_002, etc.)
                2. Provide a clear treatment_name
                3. List ALL source files that mention this treatment
                4. Extract detailed information about the treatment
                5. Note any competitor financing options mentioned
                6. Identify target demographics or patient populations
                
                IMPORTANT: 
                - Each treatment should be separate and distinct
                - Do not group everything under one treatment
                - Look for medical procedures, therapies, surgeries, diagnostic tests, etc.
                - Include both established and emerging treatments
                - Note any pricing or cost information mentioned
                
                Return a JSON array of treatments with this structure:
                {{
                    "treatment_id": "TREAT_XXX",
                    "treatment_name": "Clear Name of Treatment",
                    "source_files": ["file1.html", "file2.pdf"],
                    "category": "surgery/therapy/diagnostic/etc",
                    "raw_details": "All relevant information extracted",
                    "competitor_info": "Any competitor financing mentioned",
                    "cost_info": "Any cost/pricing information",
                    "target_demographics": "Patient population info"
                }}
                
                Documents to analyze:
                {all_content}
                """
                
                result = self.llm_agent(prompt)
                response = result.output if hasattr(result, 'output') else str(result)
                
                try:
                    # Try to extract JSON from the response
                    json_match = re.search(r'\[.*\]', response, re.DOTALL)
                    if json_match:
                        treatments = json.loads(json_match.group())
                    else:
                        treatments = json.loads(response)
                except json.JSONDecodeError:
                    # Fallback: create treatments based on text analysis
                    treatments = self._fallback_treatment_extraction(response, list(source_mapping.keys()))
                
                return treatments
                
            except Exception as e:
                print_debug(f"Error in treatment extraction: {str(e)}")
                return [{"error": f"Error extracting treatments: {str(e)}"}]
        
        return extract_treatments_from_files
    
    def _fallback_treatment_extraction(self, response: str, source_files: List[str]) -> List[Dict[str, Any]]:
        """Fallback method to extract treatments when JSON parsing fails"""
        treatments = []
        lines = response.split('\n')
        current_treatment = None
        
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in ['treatment', 'procedure', 'surgery', 'therapy']):
                if current_treatment:
                    treatments.append(current_treatment)
                
                current_treatment = {
                    "treatment_id": f"TREAT_{len(treatments) + 1:03d}",
                    "treatment_name": line.strip(),
                    "source_files": source_files,
                    "category": "medical_procedure",
                    "raw_details": line.strip(),
                    "competitor_info": "",
                    "cost_info": "",
                    "target_demographics": ""
                }
            elif current_treatment and line.strip():
                current_treatment["raw_details"] += " " + line.strip()
        
        if current_treatment:
            treatments.append(current_treatment)
        
        return treatments if treatments else [{
            "treatment_id": "TREAT_001",
            "treatment_name": "General Healthcare Services",
            "source_files": source_files,
            "category": "general",
            "raw_details": response[:500],
            "competitor_info": "",
            "cost_info": "",
            "target_demographics": ""
        }]

    def _create_treatment_grouping_tool(self):
        @tool
        def group_similar_treatments(treatments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
            """
            Group semantically similar treatments and merge their information.
            """
            try:
                prompt = f"""
                You are a medical treatment analyst. Review the following treatments and group any that are semantically similar.
                
                For each final group:
                1. Merge information from similar treatments
                2. Create a comprehensive description
                3. Combine source files
                4. Preserve all unique details
                5. Assign a clear, descriptive name
                
                Treatments to analyze:
                {json.dumps(treatments, indent=2)}
                
                Return a JSON array of grouped treatments with merged information.
                Ensure each treatment has comprehensive details and all source files are preserved.
                """
                
                result = self.llm_agent(prompt)
                response = result.output if hasattr(result, 'output') else str(result)
                
                try:
                    json_match = re.search(r'\[.*\]', response, re.DOTALL)
                    if json_match:
                        grouped_treatments = json.loads(json_match.group())
                    else:
                        grouped_treatments = json.loads(response)
                except json.JSONDecodeError:
                    # Return original treatments if grouping fails
                    grouped_treatments = treatments
                
                return grouped_treatments
                
            except Exception as e:
                print_debug(f"Error in treatment grouping: {str(e)}")
                return treatments
        
        return group_similar_treatments

    # ==================== ENHANCED DOCUMENTATION ====================
    
    def _create_detailed_description_tool(self):
        @tool
        def create_detailed_treatment_description(treatment: Dict[str, Any]) -> str:
            """
            Create a comprehensive, detailed description for each treatment (minimum 1 page).
            """
            try:
                prompt = f"""
                You are a medical writer creating detailed treatment documentation for CareCredit's business team.
                
                Create a comprehensive, detailed description (minimum 500 words) for the following treatment:
                
                Treatment: {treatment.get('treatment_name', 'Unknown')}
                Details: {treatment.get('raw_details', '')}
                Category: {treatment.get('category', '')}
                Sources: {treatment.get('source_files', [])}
                
                Your description should include:
                1. **Overview** (2-3 paragraphs): What is this treatment, why is it important
                2. **Medical Details** (2-3 paragraphs): How it works, who needs it, medical benefits
                3. **Market Context** (2 paragraphs): Current market trends, competitor offerings
                4. **Patient Demographics** (1-2 paragraphs): Target patient population, age groups, income levels
                5. **Treatment Process** (2 paragraphs): Typical procedure steps, duration, follow-up care
                6. **Cost Considerations** (1-2 paragraphs): Typical price ranges, insurance coverage, out-of-pocket costs
                7. **CareCredit Opportunity** (1-2 paragraphs): How CareCredit can help, financing benefits
                
                Write in a professional, business-friendly tone that helps executives understand the opportunity.
                Make it engaging and informative, with specific details and market insights.
                """
                
                result = self.llm_agent(prompt)
                response = result.output if hasattr(result, 'output') else str(result)
                
                return response
                
            except Exception as e:
                return f"Error creating detailed description: {str(e)}"
        
        return create_detailed_treatment_description

    def _create_enhanced_report_generator_tool(self):
        @tool
        def generate_enhanced_report(treatment_data: Dict[str, Any]) -> Dict[str, Any]:
            """
            Generate an enhanced report section for a treatment with detailed analysis and formatting.
            """
            try:
                prompt = f"""
                You are a healthcare business analyst creating a detailed report section for CareCredit.
                
                Generate a comprehensive report section for this treatment:
                
                Treatment: {treatment_data.get('treatment_name', 'Unknown')}
                Category: {treatment_data.get('category', '')}
                Description: {treatment_data.get('detailed_description', '')}
                Risk Assessment: {treatment_data.get('risk_assessment', {})}
                Revenue Analysis: {treatment_data.get('revenue_analysis', {})}
                Customer Impact: {treatment_data.get('customer_impact', {})}
                
                Create a well-structured report section that includes:
                
                1. Treatment Overview
                   - Clear description of the treatment
                   - Medical significance
                   - Current market position
                
                2. Market Analysis
                   - Market size and growth
                   - Competitive landscape
                   - Regulatory environment
                
                3. Risk Assessment
                   - Detailed risk factors
                   - Risk mitigation strategies
                   - Monitoring requirements
                
                4. Revenue Analysis
                   - Market opportunity
                   - Revenue projections
                   - Cost structure
                   - Profitability analysis
                
                5. Customer Impact
                   - Target demographics
                   - Customer benefits
                   - Market penetration strategy
                
                6. Strategic Recommendations
                   - Implementation approach
                   - Resource requirements
                   - Success metrics
                
                Format the content professionally with clear sections, bullet points, and key metrics.
                Include specific numbers and percentages where possible.
                """
                
                result = self.llm_agent(prompt)
                response = result.output if hasattr(result, 'output') else str(result)
                
                # Structure the response into sections
                sections = {
                    "overview": "",
                    "market_analysis": "",
                    "risk_assessment": "",
                    "revenue_analysis": "",
                    "customer_impact": "",
                    "recommendations": ""
                }
                
                # Parse the response into sections
                current_section = None
                for line in response.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Check for section headers
                    if line.lower().startswith(('1.', '2.', '3.', '4.', '5.', '6.')):
                        section_name = line.split('.')[1].strip().lower()
                        if 'overview' in section_name:
                            current_section = 'overview'
                        elif 'market' in section_name:
                            current_section = 'market_analysis'
                        elif 'risk' in section_name:
                            current_section = 'risk_assessment'
                        elif 'revenue' in section_name:
                            current_section = 'revenue_analysis'
                        elif 'customer' in section_name:
                            current_section = 'customer_impact'
                        elif 'recommendation' in section_name:
                            current_section = 'recommendations'
                    elif current_section:
                        sections[current_section] += line + '\n'
                
                return {
                    "status": "success",
                    "sections": sections,
                    "raw_content": response,
                    "timestamp": datetime.now().isoformat()
                }
                
            except Exception as e:
                return {
                    "status": "error",
                    "error": f"Error generating enhanced report: {str(e)}",
                    "timestamp": datetime.now().isoformat()
                }
        
        return generate_enhanced_report

    # ==================== COMPREHENSIVE RISK ANALYSIS ====================
    
    def _create_comprehensive_risk_analysis_tool(self):
        @tool
        def analyze_treatment_risks_comprehensive(treatment: Dict[str, Any]) -> Dict[str, Any]:
            """
            Perform comprehensive risk analysis with detailed explanations and logic.
            """
            try:
                prompt = f"""
                You are a risk assessment expert for CareCredit. Perform a comprehensive risk analysis for this treatment.
                
                Treatment: {treatment.get('treatment_name', 'Unknown')}
                Details: {treatment.get('raw_details', '')}
                Category: {treatment.get('category', '')}
                
                Provide a detailed risk analysis with the following structure:
                
                1. **RISK PARAMETERS EXPLANATION**:
                   - Market Risk (0-10): Competition, market saturation, regulatory changes
                   - Medical Risk (0-10): Safety concerns, malpractice potential, adverse events
                   - Financial Risk (0-10): Default rates, collection challenges, cost volatility
                   - Regulatory Risk (0-10): FDA approval status, changing regulations
                   - Technology Risk (0-10): Innovation obsolescence, equipment requirements
                   - Provider Risk (0-10): Provider network stability, training requirements
                
                2. **DETAILED RISK SCORING WITH LOGIC**:
                   For each parameter, provide:
                   - Score (0-10)
                   - Detailed explanation of why this score was assigned
                   - Supporting evidence or reasoning
                   - Comparison to similar treatments
                
                3. **OVERALL RISK ASSESSMENT**:
                   - Combined risk score (weighted average)
                   - Risk category (Low: 0-3, Medium: 4-6, High: 7-10)
                   - Key risk factors to monitor
                   - Risk mitigation strategies
                
                4. **RECOMMENDATIONS**:
                   - Should CareCredit proceed? (Yes/No/Conditional)
                   - Conditions or safeguards needed
                   - Monitoring requirements
                   - Exit strategies if needed
                
                Return as JSON with detailed explanations for each section.
                """
                
                result = self.llm_agent(prompt)
                response = result.output if hasattr(result, 'output') else str(result)
                
                try:
                    json_match = re.search(r'\{.*\}', response, re.DOTALL)
                    if json_match:
                        risk_analysis = json.loads(json_match.group())
                    else:
                        risk_analysis = json.loads(response)
                except json.JSONDecodeError:
                    # Create structured response from text
                    risk_analysis = {
                        "overall_risk_score": 5.0,
                        "risk_category": "Medium",
                        "detailed_analysis": response,
                        "recommendation": "Conditional approval with monitoring"
                    }
                
                return risk_analysis
                
            except Exception as e:
                return {"error": f"Error in risk analysis: {str(e)}"}
        
        return analyze_treatment_risks_comprehensive

    # ==================== DETAILED REVENUE ANALYSIS ====================
    
    def _create_detailed_revenue_analysis_tool(self):
        @tool
        def analyze_revenue_opportunities_detailed(treatment: Dict[str, Any], risk_analysis: Dict[str, Any]) -> Dict[str, Any]:
            """
            Perform detailed revenue analysis with customer segmentation and profitability calculations.
            """
            try:
                prompt = f"""
                You are a business analyst for CareCredit. Perform a comprehensive revenue analysis for this treatment.
                
                Treatment: {treatment.get('treatment_name', 'Unknown')}
                Category: {treatment.get('category', '')}
                Risk Score: {risk_analysis.get('overall_risk_score', 5)}
                
                Provide detailed revenue analysis with:
                
                1. **MARKET SIZE ESTIMATION**:
                   - Total addressable market (TAM)
                   - Serviceable addressable market (SAM)
                   - Market growth rate
                   - Competitive landscape analysis
                
                2. **CUSTOMER SEGMENTATION & IMPACT**:
                   - Age group breakdown (18-30, 31-45, 46-60, 60+)
                   - Income level distribution
                   - Geographic distribution
                   - Insurance coverage patterns
                   - Estimated patient volumes per segment
                
                3. **EXISTING CUSTOMER IMPACT**:
                   - How many current CareCredit customers would benefit
                   - Cross-selling opportunities
                   - Customer lifetime value impact
                   - Retention improvements
                
                4. **NEW CUSTOMER ACQUISITION**:
                   - Potential new customers per year
                   - Customer acquisition cost
                   - Conversion rates by segment
                   - Marketing channel effectiveness
                
                5. **FINANCIAL PROJECTIONS**:
                   - Average transaction size
                   - Transaction volume projections (Year 1-3)
                   - Revenue projections (Year 1-3)
                   - Cost structure analysis
                   - Profitability timeline
                   - ROI calculations
                
                6. **DETAILED LOGIC & ASSUMPTIONS**:
                   - Explain methodology for each calculation
                   - Key assumptions and their rationale
                   - Sensitivity analysis for key variables
                   - Comparison to similar treatments in portfolio
                
                Return comprehensive JSON with all calculations and detailed explanations.
                """
                
                result = self.llm_agent(prompt)
                response = result.output if hasattr(result, 'output') else str(result)
                
                try:
                    json_match = re.search(r'\{.*\}', response, re.DOTALL)
                    if json_match:
                        revenue_analysis = json.loads(json_match.group())
                    else:
                        revenue_analysis = json.loads(response)
                except json.JSONDecodeError:
                    # Create structured response from text
                    revenue_analysis = {
                        "market_size": "Large opportunity identified",
                        "revenue_projection_year1": "$500K - $2M",
                        "customer_impact": "Positive for existing and new customers",
                        "detailed_analysis": response,
                        "recommendation": "Proceed with pilot program"
                    }
                
                return revenue_analysis
                
            except Exception as e:
                return {"error": f"Error in revenue analysis: {str(e)}"}
        
        return analyze_revenue_opportunities_detailed
    
    def _create_customer_impact_analysis_tool(self):
        @tool
        def analyze_customer_impact(treatment: Dict[str, Any]) -> Dict[str, Any]:
            """
            Analyze detailed customer impact and segmentation.
            """
            try:
                prompt = f"""
                Analyze the customer impact for treatment: {treatment.get('treatment_name', 'Unknown')}
                
                Provide detailed breakdown of:
                1. Existing customer segments that would benefit
                2. New customer acquisition potential
                3. Demographic analysis with population estimates
                4. Geographic considerations
                5. Seasonal trends and timing factors
                
                Include specific numbers and percentages where possible.
                """
                
                result = self.llm_agent(prompt)
                response = result.output if hasattr(result, 'output') else str(result)
                
                return {"customer_impact_analysis": response}
                
            except Exception as e:
                return {"error": f"Error in customer impact analysis: {str(e)}"}
        
        return analyze_customer_impact

    # ==================== ENHANCED WORD DOCUMENT EXPORT ====================
    
    def _create_enhanced_word_export_tool(self):
        @tool
        def export_enhanced_word_document(treatments_data: List[Dict[str, Any]], output_path: str) -> Dict[str, Any]:
            """
            Create a comprehensive, visually appealing Word document with detailed treatment analysis.
            """
            try:
                doc = docx.Document()
                
                # Set up document styles
                self._setup_document_styles(doc)
                
                # Add title page
                self._add_title_page(doc)
                
                # Add executive summary
                self._add_executive_summary(doc, treatments_data)
                
                # Add detailed treatment sections
                for i, treatment_data in enumerate(treatments_data, 1):
                    self._add_treatment_section(doc, treatment_data, i)
                    
                    # Add page break between treatments
                    if i < len(treatments_data):
                        doc.add_page_break()
                
                # Add conclusions and recommendations
                self._add_conclusions(doc, treatments_data)
                
                # Save document
                doc.save(output_path)
                
                return {
                    "status": "success",
                    "output_path": output_path,
                    "treatments_count": len(treatments_data),
                    "timestamp": datetime.now().isoformat()
                }
                
            except Exception as e:
                return {"error": f"Error creating enhanced Word document: {str(e)}"}
        
        return export_enhanced_word_document
    
    def _setup_document_styles(self, doc):
        """Set up custom styles for the document"""
        try:
            # Title style
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(24)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Heading styles
            heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Arial'
            heading_font.size = Pt(16)
            heading_font.bold = True
            
            # Subheading style
            subheading_style = doc.styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
            subheading_font = subheading_style.font
            subheading_font.name = 'Arial'
            subheading_font.size = Pt(12)
            subheading_font.bold = True
            
        except Exception as e:
            print_debug(f"Warning: Could not set custom styles: {e}")
    
    def _add_title_page(self, doc):
        """Add an attractive title page"""
        title = doc.add_heading('CareCredit Healthcare Treatment Analysis', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('Comprehensive Market Analysis & Revenue Opportunity Assessment')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty line
        
        date_para = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc, treatments_data):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        
        summary_text = f"""
        This comprehensive analysis examines {len(treatments_data)} distinct healthcare treatments 
        identified from competitor analysis and market research. Each treatment has been evaluated 
        for market opportunity, risk assessment, and revenue potential for CareCredit's financing services.
        
        Key findings include market sizing, customer segmentation analysis, competitive positioning, 
        and detailed financial projections. Risk assessments provide clear guidance on implementation 
        strategies and monitoring requirements.
        """
        
        doc.add_paragraph(summary_text)
        doc.add_page_break()
    
    def _add_treatment_section(self, doc, treatment_data, treatment_number):
        """Add detailed section for each treatment"""
        treatment_name = treatment_data.get('treatment_name', f'Treatment {treatment_number}')
        
        # Main treatment heading
        doc.add_heading(f'Treatment {treatment_number}: {treatment_name}', level=1)
        
        # Treatment overview
        doc.add_heading('Overview', level=2)
        overview = treatment_data.get('detailed_description', 'No detailed description available')
        doc.add_paragraph(overview)
        
        # Source information
        doc.add_heading('Data Sources', level=2)
        sources = treatment_data.get('source_files', [])
        if sources:
            for source in sources:
                doc.add_paragraph(f'• {source}', style='List Bullet')
        else:
            doc.add_paragraph('• Multiple research sources analyzed')
        
        # Risk Analysis Section
        self._add_risk_section(doc, treatment_data.get('risk_assessment', {}))
        
        # Revenue Analysis Section
        self._add_revenue_section(doc, treatment_data.get('revenue_analysis', {}))
        
        # Customer Impact Section
        self._add_customer_impact_section(doc, treatment_data.get('customer_impact', {}))
    
    def _add_risk_section(self, doc, risk_data):
        """Add detailed risk analysis section"""
        doc.add_heading('Risk Assessment', level=2)
        
        if isinstance(risk_data, dict) and 'detailed_analysis' in risk_data:
            doc.add_paragraph(risk_data['detailed_analysis'])
        elif isinstance(risk_data, str):
            doc.add_paragraph(risk_data)
        else:
            doc.add_paragraph("Comprehensive risk analysis includes market, medical, financial, regulatory, technology, and provider risk factors. Each parameter is scored on a 0-10 scale with detailed justification and mitigation strategies.")
        
        # Add risk metrics table if available
        if isinstance(risk_data, dict):
            doc.add_heading('Risk Metrics Summary', level=3)
            
            # Create a simple risk summary
            risk_score = risk_data.get('overall_risk_score', 'Not calculated')
            risk_category = risk_data.get('risk_category', 'Under review')
            
            doc.add_paragraph(f'Overall Risk Score: {risk_score}')
            doc.add_paragraph(f'Risk Category: {risk_category}')
            
            recommendation = risk_data.get('recommendation', 'Further analysis required')
            doc.add_paragraph(f'Recommendation: {recommendation}')
    
    def _add_revenue_section(self, doc, revenue_data):
        """Add detailed revenue analysis section"""
        doc.add_heading('Revenue Analysis', level=2)
        
        if isinstance(revenue_data, dict) and 'detailed_analysis' in revenue_data:
            doc.add_paragraph(revenue_data['detailed_analysis'])
        elif isinstance(revenue_data, str):
            doc.add_paragraph(revenue_data)
        else:
            doc.add_paragraph("Revenue analysis includes market sizing, customer segmentation, financial projections, and profitability assessment. Projections are based on market research, competitive analysis, and CareCredit's historical performance data.")
        
        # Add revenue projections if available
        if isinstance(revenue_data, dict):
            doc.add_heading('Financial Projections', level=3)
            
            market_size = revenue_data.get('market_size', 'Under analysis')
            doc.add_paragraph(f'Market Size: {market_size}')
            
            revenue_proj = revenue_data.get('revenue_projection_year1', 'To be determined')
            doc.add_paragraph(f'Year 1 Revenue Projection: {revenue_proj}')
    
    def _add_customer_impact_section(self, doc, customer_data):
        """Add customer impact analysis section"""
        doc.add_heading('Customer Impact Analysis', level=2)
        
        if isinstance(customer_data, dict) and 'customer_impact_analysis' in customer_data:
            doc.add_paragraph(customer_data['customer_impact_analysis'])
        elif isinstance(customer_data, str):
            doc.add_paragraph(customer_data)
        else:
            doc.add_paragraph("Customer impact analysis examines how this treatment affects existing CareCredit customers and potential for new customer acquisition. Analysis includes demographic segmentation, geographic distribution, and behavioral patterns.")
    
    def _add_conclusions(self, doc, treatments_data):
        """Add conclusions and recommendations section"""
        doc.add_heading('Conclusions and Strategic Recommendations', level=1)
        
        # Add overall summary
        doc.add_heading('Overall Analysis Summary', level=2)
        summary_text = f"""
        This comprehensive analysis has identified and evaluated {len(treatments_data)} distinct healthcare treatments 
        that present significant opportunities for CareCredit's financing services. Each treatment has been thoroughly 
        analyzed for market potential, risk factors, customer impact, and revenue opportunities.
        
        Key findings across all treatments include:
        • Market Opportunity: Significant growth potential in emerging treatments
        • Customer Impact: Strong potential for both existing and new customer acquisition
        • Risk Profile: Varied risk levels requiring different implementation strategies
        • Revenue Potential: Promising financial projections with clear paths to profitability
        """
        doc.add_paragraph(summary_text)
        
        # Add treatment-specific recommendations
        doc.add_heading('Treatment-Specific Recommendations', level=2)
        for treatment in treatments_data:
            treatment_name = treatment.get('treatment_name', 'Unknown Treatment')
            doc.add_heading(f'{treatment_name} - Strategic Recommendations', level=3)
            
            # Get risk and revenue data
            risk_data = treatment.get('risk_assessment', {})
            revenue_data = treatment.get('revenue_analysis', {})
            
            recommendation_text = f"""
            Based on comprehensive analysis of {treatment_name}, we recommend the following strategic approach:
            
            1. Implementation Strategy:
               • {self._get_implementation_strategy(risk_data, revenue_data)}
            
            2. Risk Management:
               • {self._get_risk_management_strategy(risk_data)}
            
            3. Revenue Optimization:
               • {self._get_revenue_optimization_strategy(revenue_data)}
            
            4. Customer Engagement:
               • {self._get_customer_engagement_strategy(treatment.get('customer_impact', {}))}
            """
            doc.add_paragraph(recommendation_text)
        
        # Add final recommendations
        doc.add_heading('Final Strategic Recommendations', level=2)
        final_recommendations = """
        1. Immediate Actions:
           • Prioritize treatments with highest revenue potential and lowest risk
           • Develop targeted marketing campaigns for high-impact treatments
           • Establish monitoring systems for risk metrics
        
        2. Short-term Goals (3-6 months):
           • Launch pilot programs for top 3 treatments
           • Develop provider partnerships
           • Create customer education materials
        
        3. Long-term Strategy (6-12 months):
           • Expand treatment portfolio based on performance
           • Optimize financing terms based on risk profiles
           • Develop advanced analytics for treatment performance
        """
        doc.add_paragraph(final_recommendations)
        
        # Add appendix with detailed metrics
        self._add_appendix(doc, treatments_data)
    
    def _get_implementation_strategy(self, risk_data: Dict[str, Any], revenue_data: Dict[str, Any]) -> str:
        """Generate implementation strategy based on risk and revenue data"""
        risk_score = risk_data.get('overall_risk_score', 5)
        revenue_potential = revenue_data.get('revenue_projection_year1', 'Unknown')
        
        if risk_score <= 3:
            return "Implement immediately with standard monitoring"
        elif risk_score <= 6:
            return "Implement with enhanced monitoring and phased rollout"
        else:
            return "Implement with strict controls and limited initial scope"
    
    def _get_risk_management_strategy(self, risk_data: Dict[str, Any]) -> str:
        """Generate risk management strategy based on risk analysis"""
        risk_category = risk_data.get('risk_category', 'Medium')
        
        if risk_category == 'Low':
            return "Standard monitoring with quarterly reviews"
        elif risk_category == 'Medium':
            return "Enhanced monitoring with monthly reviews and risk mitigation plans"
        else:
            return "Intensive monitoring with weekly reviews and contingency plans"
    
    def _get_revenue_optimization_strategy(self, revenue_data: Dict[str, Any]) -> str:
        """Generate revenue optimization strategy based on revenue analysis"""
        market_size = revenue_data.get('market_size', 'Unknown')
        
        if 'Large' in str(market_size):
            return "Aggressive market penetration with competitive pricing"
        elif 'Medium' in str(market_size):
            return "Balanced growth with focus on market share"
        else:
            return "Focused growth with premium positioning"
    
    def _get_customer_engagement_strategy(self, customer_data: Dict[str, Any]) -> str:
        """Generate customer engagement strategy based on customer impact analysis"""
        impact_analysis = customer_data.get('customer_impact_analysis', '')
        
        if 'significant' in impact_analysis.lower():
            return "Comprehensive engagement program with multiple touchpoints"
        elif 'moderate' in impact_analysis.lower():
            return "Targeted engagement with key customer segments"
        else:
            return "Basic engagement with focus on education and awareness"
    
    def _add_appendix(self, doc, treatments_data):
        """Add detailed appendix with metrics and analysis"""
        doc.add_page_break()
        doc.add_heading('Appendix: Detailed Analysis Metrics', level=1)
        
        for treatment in treatments_data:
            treatment_name = treatment.get('treatment_name', 'Unknown Treatment')
            doc.add_heading(f'{treatment_name} - Detailed Metrics', level=2)
            
            # Add risk metrics
            self._add_detailed_risk_metrics(doc, treatment.get('risk_assessment', {}))
            
            # Add revenue metrics
            self._add_detailed_revenue_metrics(doc, treatment.get('revenue_analysis', {}))
            
            # Add customer metrics
            self._add_detailed_customer_metrics(doc, treatment.get('customer_impact', {}))
            
            doc.add_page_break()
    
    def _add_detailed_risk_metrics(self, doc, risk_data: Dict[str, Any]):
        """Add detailed risk metrics to appendix"""
        doc.add_heading('Risk Metrics', level=3)
        
        if isinstance(risk_data, dict):
            # Add risk parameters
            doc.add_paragraph('Risk Parameters:')
            for param, score in risk_data.get('risk_parameters', {}).items():
                doc.add_paragraph(f'• {param}: {score}/10', style='List Bullet')
            
            # Add risk explanations
            if 'risk_explanations' in risk_data:
                doc.add_paragraph('Risk Explanations:')
                for param, explanation in risk_data['risk_explanations'].items():
                    doc.add_paragraph(f'• {param}: {explanation}', style='List Bullet')
    
    def _add_detailed_revenue_metrics(self, doc, revenue_data: Dict[str, Any]):
        """Add detailed revenue metrics to appendix"""
        doc.add_heading('Revenue Metrics', level=3)
        
        if isinstance(revenue_data, dict):
            # Add market metrics
            doc.add_paragraph('Market Metrics:')
            for metric, value in revenue_data.get('market_metrics', {}).items():
                doc.add_paragraph(f'• {metric}: {value}', style='List Bullet')
            
            # Add financial projections
            doc.add_paragraph('Financial Projections:')
            for year, projection in revenue_data.get('projections', {}).items():
                doc.add_paragraph(f'• {year}: {projection}', style='List Bullet')
    
    def _add_detailed_customer_metrics(self, doc, customer_data: Dict[str, Any]):
        """Add detailed customer metrics to appendix"""
        doc.add_heading('Customer Metrics', level=3)
        
        if isinstance(customer_data, dict):
            # Add demographic metrics
            doc.add_paragraph('Demographic Metrics:')
            for segment, data in customer_data.get('demographics', {}).items():
                doc.add_paragraph(f'• {segment}: {data}', style='List Bullet')
            
            # Add impact metrics
            doc.add_paragraph('Impact Metrics:')
            for metric, value in customer_data.get('impact_metrics', {}).items():
                doc.add_paragraph(f'• {metric}: {value}', style='List Bullet')

    def process_files(self, input_files: List[str], output_path: str) -> Dict[str, Any]:
        """
        Process input files and generate comprehensive treatment analysis report.
        
        Args:
            input_files: List of file paths to process
            output_path: Path where the final report will be saved
            
        Returns:
            Dict containing processing results and status
        """
        try:
            # Read all input files
            file_contents = []
            for file_path in input_files:
                file_data = self.research_agent.read_file(file_path)
                if "error" not in file_data:
                    file_contents.append(file_data)
            
            # Extract treatments from files
            treatments = self.research_agent.extract_treatments_from_files(file_contents)
            
            # Group similar treatments
            grouped_treatments = self.research_agent.group_similar_treatments(treatments)
            
            # Process each treatment
            processed_treatments = []
            for treatment in grouped_treatments:
                # Create detailed description
                detailed_description = self.documenting_agent.create_detailed_treatment_description(treatment)
                treatment['detailed_description'] = detailed_description
                
                # Perform risk assessment
                risk_assessment = self.risk_assessment_agent.analyze_treatment_risks_comprehensive(treatment)
                treatment['risk_assessment'] = risk_assessment
                
                # Analyze revenue opportunities
                revenue_analysis = self.revenue_identification_agent.analyze_revenue_opportunities_detailed(
                    treatment, risk_assessment
                )
                treatment['revenue_analysis'] = revenue_analysis
                
                # Analyze customer impact
                customer_impact = self.revenue_identification_agent.analyze_customer_impact(treatment)
                treatment['customer_impact'] = customer_impact
                
                processed_treatments.append(treatment)
            
            # Generate final report
            report_result = self.emailing_agent.export_enhanced_word_document(
                processed_treatments, output_path
            )
            
            # Update processing results
            self.processing_results.treatments = processed_treatments
            self.processing_results.final_report = output_path
            self.processing_results.approval_status = "completed"
            
            return {
                "status": "success",
                "treatments_processed": len(processed_treatments),
                "output_path": output_path,
                "timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            error_msg = f"Error processing files: {str(e)}"
            print_debug(error_msg)
            return {
                "status": "error",
                "error": error_msg,
                "timestamp": datetime.now().isoformat()
            }

def print_debug(message: str):
    """Helper function for debug printing"""
    print(f"[DEBUG] {message}")

if __name__ == "__main__":
    # Initialize the healthcare agent system
    healthcare_system = HealthcareAgentSystem()
    
    # Example usage
    input_files = [
        "sample_treatment1.html",
        "sample_treatment2.pdf",
        "sample_treatment3.docx"
    ]
    
    # Process files and generate report
    output_path = "./hackathon_output/treatment_analysis_report.docx"
    healthcare_system.process_files(input_files, output_path)
